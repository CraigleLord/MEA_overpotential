"""
Convert speaker script to a formatted Word document.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

BASE = r"c:\Users\user\My Drive\KAIST MASc 2021\Laboratory Work\Protocol\overpotential calculation\For paper SI"
OUT  = BASE + r"\Te_Pt3Co_speaker_script.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0D, 0x2B, 0x55)
TEAL  = RGBColor(0x00, 0x7A, 0x87)
GRAY  = RGBColor(0x55, 0x55, 0x55)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_rule(doc, color_hex="0D2B55"):
    """Horizontal rule via bottom border on a blank paragraph."""
    p = doc.add_paragraph()
    para_spacing(p, before=2, after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ── Document title block ──────────────────────────────────────────────────────
p = doc.add_paragraph()
para_spacing(p, before=0, after=4)
run = p.add_run("Speaker Script — Literature Talk")
set_font(run, size=20, bold=True, color=NAVY)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
para_spacing(p, before=0, after=2)
run = p.add_run(
    "“Thermostable Tellurium Anchoring Enabling Robust Thermal and\n"
    "Electrochemical Stability for Pt₃Co Intermetallic Fuel Cell Catalysts”"
)
set_font(run, size=13, bold=True, italic=True, color=TEAL)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
para_spacing(p, before=0, after=2)
run = p.add_run("Chen et al., Advanced Functional Materials 2024")
set_font(run, size=11, italic=True, color=GRAY)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
para_spacing(p, before=2, after=6)
run = p.add_run("Total estimated time: ~17–19 minutes")
set_font(run, size=11, bold=True, color=NAVY)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_rule(doc)

# ── Slide data ────────────────────────────────────────────────────────────────
slides = [
    {
        "header": "SLIDE 1 — Title",
        "timing": "0:00 – 0:45  (∼45 sec)",
        "body": (
            "Good morning/afternoon everyone. Today I’ll be presenting a paper from the journal "
            "Advanced Functional Materials, published in 2024 by Yuanxin Chen and colleagues from "
            "Wuhan University of Technology.\n\n"
            "The paper is titled: “Thermostable Tellurium Anchoring Enabling Robust Thermal and "
            "Electrochemical Stability for Pt₃Co Intermetallic Fuel Cell Catalysts.”\n\n"
            "At its heart, this work is about solving one of the most persistent challenges in PEMFC "
            "catalyst design — getting ordered intermetallic nanoparticles to stay small, active, "
            "and stable. And their solution is surprisingly elegant: use tellurium.\n\n"
            "Let me walk you through the whole story."
        ),
    },
    {
        "header": "SLIDE 2 — Background: Why Better PEMFC Catalysts?",
        "timing": "0:45 – 2:30  (∼1 min 45 sec)",
        "body": (
            "To set the stage: proton-exchange membrane fuel cells, PEMFCs, convert hydrogen gas into "
            "electricity. The only byproduct is water, which makes them a clean energy conversion "
            "technology. They’re already being deployed in hydrogen buses, forklifts, and automotive "
            "applications like the Toyota Mirai.\n\n"
            "The critical limiting factor is the cathode, where the oxygen reduction reaction — ORR "
            "— takes place. This reaction is sluggish. It needs a good catalyst to proceed efficiently "
            "at practical rates.\n\n"
            "The current standard is platinum nanoparticles on carbon black. But platinum is expensive "
            "and rare, and the nanoparticles degrade over time through dissolution, Ostwald ripening, "
            "and carbon corrosion.\n\n"
            "To reduce Pt usage and improve activity, researchers alloy Pt with cheaper 3d transition "
            "metals like cobalt. Ordered intermetallic Pt₃Co — with the L1₂ crystal structure "
            "— is particularly attractive. The periodic atomic arrangement creates beneficial strain "
            "and ligand effects that boost intrinsic ORR activity and thermodynamic stability.\n\n"
            "The US Department of Energy has set 2025 targets: mass activity of at least 0.44 A per "
            "milligram of Pt at 0.9 V, and at least 60% retention after 30,000 accelerated voltage "
            "cycles. These are the benchmarks this paper is measured against."
        ),
    },
    {
        "header": "SLIDE 3 — The Problem: Sintering During High-Temperature Annealing",
        "timing": "2:30 – 4:15  (∼1 min 45 sec)",
        "body": (
            "So why isn’t Pt₃Co already in every fuel cell? The problem lies in its synthesis.\n\n"
            "To form the ordered intermetallic phase, you need high-temperature annealing — typically "
            "above 700°C — to give atoms enough energy to diffuse and rearrange into the ordered L1₂ "
            "lattice. This is non-negotiable.\n\n"
            "But here’s the catch: that same high temperature causes the nanoparticles to sinter. On "
            "bare carbon support, Pt₃Co particles grow from a desirable 3–4 nm to 9 nm or more. Why "
            "does this matter? Because ECSA — electrochemical surface area — scales with particle "
            "size. Larger particles mean fewer surface Pt atoms per gram, meaning lower mass activity.\n\n"
            "Sintering proceeds through two pathways: particle migration and coalescence, where whole "
            "particles drift and merge; and Ostwald ripening, where Pt atoms dissolve off smaller "
            "particles and redeposit on larger ones.\n\n"
            "The root cause is weak metal-support interactions between Pt and bare carbon. The particles "
            "aren’t anchored; they’re essentially free to wander at high temperatures.\n\n"
            "Now, people have tried various fixes — doping carbon with sulfur, adding phosphorus-oxygen "
            "groups. These help, but only partially. The key insight in this paper is that sulfur and "
            "selenium — the most common heteroatom dopants — actually evaporate during high-temperature "
            "annealing. So even if they help initially, they’re gone by the time the catalyst is done. "
            "Tellurium, being a much heavier semimetal, stays."
        ),
    },
    {
        "header": "SLIDE 4 — Strategy: Te-C Support by CVD",
        "timing": "4:15 – 5:30  (∼1 min 15 sec)",
        "body": (
            "Here’s the overall strategy. This is the schematic in Figure 1a.\n\n"
            "The authors use chemical vapour deposition — a simple, industrially relevant technique "
            "— to deposit tellurium onto commercial Ketjen Black carbon (EC-300J). This creates the "
            "Te-C support. No toxic organic solvents, relatively low energy process.\n\n"
            "Then, the standard wet impregnation protocol deposits Pt and Co precursors onto Te-C. "
            "This is followed by a two-stage heat treatment: a high-temperature step to alloy Pt and "
            "Co and create a disordered intermetallic structure, then a lower-temperature holding step "
            "to drive the disorder-to-order transition and form the L1₂ phase.\n\n"
            "For comparison, they synthesise three other catalysts in parallel: Pt₃Co on Se-doped "
            "carbon, S-doped carbon, and undoped carbon. This gives a clean four-way comparison with "
            "the same protocol, same Pt/Co loading, only the support chemistry differs."
        ),
    },
    {
        "header": "SLIDE 5 — Structural Characterisation: Particle Size and Ordering",
        "timing": "5:30 – 7:00  (∼1 min 30 sec)",
        "body": (
            "The first thing they measure is particle size and ordering, using XRD and HAADF-STEM.\n\n"
            "The XRD pattern is highly diagnostic. For Pt₃Co, the characteristic superlattice peaks "
            "— specifically the (110) peak — only appear in a fully ordered L1₂ phase. In "
            "Pt₃Co/C, that peak is absent. In Pt₃Co/Te-C, it’s clearly present. And the "
            "I₁₁₀/I₁₁₁ ratio — a quantitative measure of ordering — is highest for "
            "Te-C at 0.212, compared to 0.105 for bare carbon. Te doping doesn’t just prevent "
            "sintering; it actively promotes ordering.\n\n"
            "From the Scherrer equation applied to the (111) peak width, they extract particle sizes: "
            "3.9 nm for Te-C, versus 9.1 nm for bare carbon. That’s a factor of 2.3 in linear size, "
            "which translates to a factor of 12 in volume.\n\n"
            "HAADF-STEM images confirm this visually — the Te-C sample shows uniformly distributed "
            "small nanoparticles, while the bare carbon sample shows dramatically larger, unevenly "
            "distributed particles with wide size distributions."
        ),
    },
    {
        "header": "SLIDE 6 — Atomic-Resolution STEM",
        "timing": "7:00 – 8:00  (∼1 min)",
        "body": (
            "For the Pt₃Co/Te-C sample, they go further with atomic-resolution HAADF-STEM imaging.\n\n"
            "This image shows a single 4 nm Pt₃Co nanoparticle viewed along the [001] zone axis. In "
            "HAADF imaging, intensity scales with atomic number — so the brighter spots are Pt columns "
            "and the dimmer spots are Co columns. You can see them alternating in a perfect square "
            "lattice — that’s the L1₂ structure.\n\n"
            "The FFT pattern shows the expected superlattice spots for the ordered phase. The line "
            "intensity profile across the particle shows the alternating contrast of Pt and Co columns.\n\n"
            "The EDS mapping confirms that Pt and Co are uniformly distributed through the particle "
            "with the correct 3:1 ratio. And critically — you can see residual Te atoms enriched in "
            "the regions adjacent to the nanoparticles. This is the first visual evidence of where "
            "the Te actually sits and why it’s effective: it’s right at the metal-support interface."
        ),
    },
    {
        "header": "SLIDE 7 — In-situ Thermal Stability",
        "timing": "8:00 – 9:00  (∼1 min)",
        "body": (
            "To see the sintering process in real time, the authors performed in-situ high-temperature "
            "XRD. They heated the precatalyst samples from room temperature to 700°C and held for 6 "
            "hours while continuously collecting XRD patterns.\n\n"
            "Figure 2 shows the evolution. As temperature rises, the (111) peak sharpens — telling us "
            "particles are growing. But the rate is dramatically different.\n\n"
            "In Pt₃Co/C, the particle size reaches nearly 7 nm after 6 hours at 700°C. In "
            "Pt₃Co/Te-C, it reaches only 4.1 nm — barely growing at all. The Te-C support is "
            "effectively pinning the particles in place at temperatures where bare carbon provides "
            "no resistance to sintering.\n\n"
            "There’s another key observation: the (110) superlattice peak — the marker of the ordered "
            "phase — appears in Pt₃Co/Te-C during the heating but never appears in Pt₃Co/C. This "
            "confirms that the same Te anchoring that suppresses sintering also facilitates ordering. "
            "These two beneficial effects appear to be linked through the stronger metal-support interaction."
        ),
    },
    {
        "header": "SLIDE 8 — Electronic Structure: XPS",
        "timing": "9:00 – 10:00  (∼1 min)",
        "body": (
            "How does Te actually interact with the Pt₃Co nanoparticles? The answer comes from XPS.\n\n"
            "Looking at the Pt 4f spectra: in Pt₃Co/C, the Pt 4f₇/₂ peak sits at 71.54 eV. In "
            "Pt₃Co/Te-C, it shifts to higher binding energy by 0.12 eV. This is a positive shift, "
            "indicating electron transfer from Pt to the Te-doped carbon support. In other words, "
            "Te acts as an electron acceptor, pulling electron density from Pt.\n\n"
            "This has two consequences. First, it strengthens the Pt–support interaction — the Pt "
            "is essentially electrostatically anchored. Second, it downshifts the d-band centre of "
            "Pt, which according to the d-band model weakens the adsorption of oxygen-containing "
            "intermediates on the Pt surface. Weaker *OH binding means easier *OH desorption, which "
            "as we’ll see is the rate-limiting step for ORR.\n\n"
            "The Co 2p shift is even more dramatic — 0.95 eV positive shift in Pt₃Co/Te-C. This "
            "suggests that Co is also involved in electron transfer to Te-C, which may explain why "
            "the L1₂ ordering is enhanced: the Te pulls electron density from Co, stabilising Co "
            "in the lattice and preventing its preferential surface segregation or dissolution."
        ),
    },
    {
        "header": "SLIDE 9 — Coordination Structure: XAS/EXAFS",
        "timing": "10:00 – 11:15  (∼1 min 15 sec)",
        "body": (
            "XPS tells us about oxidation states and electron transfer. X-ray absorption spectroscopy "
            "gives us the coordination environment.\n\n"
            "The XANES at the Pt L₃ edge shows Pt in Pt₃Co/Te-C is essentially metallic — the white "
            "line intensity matches Pt foil, and quantitative analysis gives a Pt valence of +0.13. "
            "So despite the electron transfer seen in XPS, Pt remains reduced in the bulk.\n\n"
            "The EXAFS is where it gets interesting. In R-space, Pt₃Co/Te-C shows a broad peak "
            "centred around 2.2 Å. Fitting reveals three overlapping shells:\n"
            "    •  Pt–Pt at 2.60 Å, coordination number 7.7\n"
            "    •  Pt–Co at 2.68 Å, CN 2.2\n"
            "    •  Pt–Te at 2.65 Å, CN 1.3  ← direct proof of bonding to the support\n\n"
            "That Pt–Te shell is the direct proof of bonding between the nanoparticle and the "
            "Te-doped support. The coordination number of 1.3 per Pt atom means a significant "
            "fraction of surface Pt atoms are directly bonded to Te in the support.\n\n"
            "Also note: all three bond lengths are shorter than Pt foil (2.425 Å). This compressive "
            "strain is consistent with the ligand and strain effects expected to enhance ORR "
            "intrinsic activity."
        ),
    },
    {
        "header": "SLIDE 10 — Thermostability of Te",
        "timing": "11:15 – 12:00  (∼45 sec)",
        "body": (
            "This slide directly addresses the ‘why Te?’ question.\n\n"
            "After the full high-temperature annealing cycle, the authors measured how much heteroatom "
            "remained using ICP-OES. The results are stark:\n"
            "    •  Te:  2.3 wt% retained — clearly visible in XPS\n"
            "    •  Se:  0.9 wt% — barely detectable\n"
            "    •  S:   0.2 wt% — essentially gone; no XPS signal\n\n"
            "This is the crucial result. S and Se — the popular heteroatom dopants in the literature "
            "— largely evaporate during annealing. Te, being a much heavier element with lower vapour "
            "pressure, stays on the support. And because it stays, the MSI-enhancing Pt–Te bonds are "
            "maintained not just during synthesis but throughout the entire operational life of the "
            "fuel cell.\n\n"
            "This is the mechanistic logic that ties the whole paper together."
        ),
    },
    {
        "header": "SLIDE 11 — ORR Activity in RDE",
        "timing": "12:00 – 13:00  (∼1 min)",
        "body": (
            "Now for the electrochemical results. In rotating disk electrode tests at room temperature "
            "in 0.1 M perchloric acid:\n\n"
            "The half-wave potential of Pt₃Co/Te-C is 0.925 V — 30 mV higher than Pt₃Co/C at "
            "0.895 V. For ORR, every millivolt of E₁/₂ improvement represents a real performance gain.\n\n"
            "At 0.9 V, the specific activity of Pt₃Co/Te-C is 0.64 mA/cm²Pt and mass activity is "
            "301 mA/mgPt. Compared to Pt₃Co/C at 0.41 mA/cm²Pt SA and 103 mA/mgPt MA, that’s "
            "1.56× and 2.92× improvements respectively.\n\n"
            "The ECSA of Pt₃Co/Te-C is 47.3 m²/gPt, versus only 25.2 m²/gPt for Pt₃Co/C. This "
            "factor-of-two difference in ECSA directly reflects the factor-of-two difference in "
            "particle size. More surface area per gram of Pt means more active sites per milligram "
            "— exactly why keeping particles small matters so much."
        ),
    },
    {
        "header": "SLIDE 12 — Durability in RDE: 100,000 Cycles",
        "timing": "13:00 – 14:00  (∼1 min)",
        "body": (
            "The durability results are where this paper really distinguishes itself.\n\n"
            "Pt₃Co/Te-C was subjected to 100,000 accelerated voltage cycles between 0.6 and 1.0 V. "
            "Let that number sink in — 100,000 cycles. Most literature reports 10,000 to 30,000.\n\n"
            "After those 100,000 cycles: E₁/₂ lost only 8 mV. Specific activity loss: just 1.5%. "
            "The TEM images before and after the ADT look essentially identical — no Ostwald "
            "ripening, no coalescence, no detachment.\n\n"
            "By contrast, Pt₃Co/C showed 20 mV E₁/₂ loss and 5% SA loss after only 30,000 cycles "
            "— and TEM showed clear agglomeration and loss of particle count.\n\n"
            "The 1.5% SA retention in Pt₃Co/Te-C is, by the authors’ own literature comparison, "
            "among the best ever reported for Pt-based catalysts. The Te–Pt bonding is so effective "
            "that the electrochemical operating conditions that would normally dissolve, migrate, and "
            "detach Pt nanoparticles simply cannot overcome the anchoring force."
        ),
    },
    {
        "header": "SLIDE 13 — MEA Performance: Exceeds DOE 2025 Targets",
        "timing": "14:00 – 15:15  (∼1 min 15 sec)",
        "body": (
            "RDE is a useful screening tool, but the real test is in a membrane electrode assembly "
            "— the actual fuel cell device. The authors built MEAs with Pt₃Co/Te-C and Pt₃Co/C as "
            "cathodes, with a Pt loading of 0.2 mg/cm² matching the NEDO 2030 target.\n\n"
            "At beginning of life under H₂–O₂ at 80°C:\n"
            "    •  Pt₃Co/Te-C mass activity at 0.9 V: 0.50 A/mgPt — exceeds DOE target of 0.44\n"
            "    •  Peak power density: 2.32 W/cm² at 4 A/cm² — impressively high\n"
            "    •  Pt₃Co/C by comparison: only 0.165 A/mgPt — a factor of 3 lower\n\n"
            "After 30,000 DOE protocol voltage cycles:\n"
            "    •  Pt₃Co/Te-C MA retention: 74% — exceeds DOE target of 60%\n"
            "    •  Pt₃Co/C MA attenuation: 72.7% lost — essentially failed\n\n"
            "What’s remarkable here is not just that Pt₃Co/Te-C exceeds both DOE targets, but that "
            "it does so at a platinum loading already aligned with commercial viability. This is not "
            "just a laboratory curiosity — it’s pointing at a real device-level solution."
        ),
    },
    {
        "header": "SLIDE 14 — Post-Durability MEA Analysis",
        "timing": "15:15 – 16:00  (∼45 sec)",
        "body": (
            "To understand what happened inside the MEA during the durability test, the authors used "
            "small-angle X-ray scattering — SAXS — and HAADF-STEM on the end-of-life electrode.\n\n"
            "SAXS shows that the predominant particle size shifted from 4 nm at BOL to about 3 nm at "
            "EOL — actually getting slightly smaller. This is consistent with Pt surface dissolution "
            "through Ostwald ripening during the high-potential excursions. Some Pt is lost, but the "
            "remaining particles maintain their structural integrity.\n\n"
            "The key result is in the atomic-resolution STEM at EOL: Pt₃Co on Te-C still shows the "
            "L1₂ ordered structure. The FFT patterns confirm it. On bare carbon at EOL, the ordered "
            "structure is gone — replaced by disordered alloy. This is because bare carbon couldn’t "
            "prevent Co dissolution during the durability test, and once Co content drops, the ordered "
            "phase destabilises. Te-C keeps Co inside the nanoparticle, preserving the ordered phase "
            "that’s responsible for the high intrinsic activity."
        ),
    },
    {
        "header": "SLIDE 15 — DFT Calculations",
        "timing": "16:00 – 17:00  (∼1 min)",
        "body": (
            "The authors complement the experiments with density functional theory calculations to "
            "provide mechanistic understanding.\n\n"
            "First, they modelled a 13-atom Pt₁₀Co₃ cluster adsorbed on graphene with and without "
            "heteroatom dopants. The calculated binding energies are:\n"
            "    •  Te-C:   −10.28 eV  (strongest MSI)\n"
            "    •  Se-C:   −9.49 eV\n"
            "    •  S-C:    −8.61 eV\n"
            "    •  Bare C: −3.95 eV  (weakest)\n\n"
            "The more negative the value, the stronger the metal-support interaction. Te-C wins by a "
            "wide margin. The differential charge density maps show electron density accumulating at "
            "the metal-support interface with Te doping — direct confirmation of the charge transfer "
            "mechanism seen in XPS.\n\n"
            "For ORR activity, they modelled the four-electron pathway on Pt(111) surfaces. The "
            "rate-determining step in all cases is *OH desorption. The Te-C system has the smallest "
            "free energy barrier for this step, with a limiting potential of 0.64 V — the best among "
            "all four. This explains the highest experimental ORR activity of Pt₃Co/Te-C."
        ),
    },
    {
        "header": "SLIDE 16 — Conclusions",
        "timing": "17:00 – 18:00  (∼1 min)",
        "body": (
            "To summarise the key contributions of this paper:\n\n"
            "1.  A thermostable Te-anchoring strategy using CVD is demonstrated to solve the "
            "long-standing sintering problem in Pt₃Co intermetallic synthesis.\n\n"
            "2.  The mechanism is clearly identified — Pt–Te bonds at the metal-support interface "
            "create strong MSIs that prevent particle migration and promote ordered phase formation.\n\n"
            "3.  The resulting catalyst, Pt₃Co/Te-C, achieves some of the best durability numbers "
            "in the literature: only 1.5% SA loss after 100,000 RDE cycles.\n\n"
            "4.  MEA tests show both beginning-of-life performance and durability retention exceed "
            "DOE 2025 targets at a practical Pt loading.\n\n"
            "5.  DFT calculations quantitatively confirm both the stronger MSI and the improved ORR "
            "kinetics from Te doping.\n\n"
            "The central takeaway: the thermostability of Te — something intrinsic to its chemistry "
            "as a heavy semimetal — is what makes all of this work."
        ),
    },
    {
        "header": "SLIDE 17 — Discussion & Outlook",
        "timing": "18:00 – 19:00  (∼1 min)",
        "body": (
            "I’ll close with a few discussion points worth thinking about.\n\n"
            "What makes this paper particularly strong is the mechanistic coherence. Every piece of "
            "characterisation — XPS, EXAFS, in-situ XRD, DFT — tells the same consistent story. "
            "You don’t often see that level of alignment between theory and experiment.\n\n"
            "Some questions worth thinking about or discussing:\n\n"
            "The MEA tests are in H₂–O₂, not H₂–air. Air contains nitrogen, and at lower O₂ "
            "partial pressure, performance drops significantly. It would be important to see how "
            "Pt₃Co/Te-C performs in real air conditions before claiming full DOE target compliance.\n\n"
            "The paper doesn’t address the cost of Te at scale. Tellurium is a genuine critical "
            "mineral — it’s used in solar cells (CdTe), and its supply is limited. Whether Te-doped "
            "carbon is cost-effective at automotive production volumes is an open question.\n\n"
            "And finally: can this principle generalise? The idea of choosing a heteroatom dopant "
            "based on its thermostability under your synthesis conditions seems broadly applicable "
            "— to PtNi, PtFe, or other ordered intermetallic catalysts. That’s perhaps the most "
            "broadly useful insight from this work.\n\n"
            "Thank you — happy to take any questions."
        ),
    },
]

# ── Write each slide block ────────────────────────────────────────────────────
for slide in slides:
    # Slide header (navy, bold)
    p = doc.add_paragraph()
    para_spacing(p, before=14, after=2)
    run = p.add_run(slide["header"])
    set_font(run, size=14, bold=True, color=NAVY)

    # Timing badge (teal, italic, smaller)
    p2 = doc.add_paragraph()
    para_spacing(p2, before=0, after=4)
    run2 = p2.add_run(slide["timing"])
    set_font(run2, size=10, italic=True, color=TEAL)

    # Body text
    for chunk in slide["body"].split("\n\n"):
        p3 = doc.add_paragraph()
        para_spacing(p3, before=0, after=5, line=14)
        run3 = p3.add_run(chunk)
        set_font(run3, size=11, color=RGBColor(0x22, 0x22, 0x22))

    add_rule(doc, "AAAAAA")

# ── Footer note ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
para_spacing(p, before=6, after=0)
run = p.add_run(
    "End of script.  Total: ~17–19 minutes at a comfortable speaking pace.  "
    "Allow 5–10 minutes additional for Q&A."
)
set_font(run, size=10, italic=True, color=GRAY)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(f"Saved: {OUT}")
